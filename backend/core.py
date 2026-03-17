from __future__ import annotations

import html
import io
import json
import math
import os
import re
import zipfile
import xml.etree.ElementTree as ET
from collections import defaultdict
from dataclasses import dataclass
from functools import lru_cache
from typing import Dict, List, Optional, Sequence, Set, Tuple

import cv2
import fitz
import numpy as np
import requests
from docx import Document
from natasha import Doc, NewsEmbedding, Segmenter

try:
    import easyocr

    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

try:
    from natasha import NewsNERTagger as NewsNerTagger
except Exception:
    from natasha import NewsNerTagger as NewsNerTagger


SUPPORTED_EXTENSIONS = {"pdf", "docx", "txt", "md", "png", "jpg", "jpeg", "zip"}

RE_EMAIL = re.compile(
    r"\b[A-Za-z0-9._%+\-]+\s*@\s*[A-Za-z0-9.\-]+(?:\s*\.\s*|\s+)[A-Za-z]{2,}\b",
    re.IGNORECASE,
)
RE_PHONE = re.compile(
    r"(?<!\d)(?:"
    r"(?:\+7|8)(?:[\s\-\(\)]*[0-9OoОо]){10,11}"
    r"|"
    r"\(?\d{3}\)?[\s\-]+\d{3}[\s\-]*\d{2}[\s\-]*\d{2}"
    r"|"
    r"\+7[\s\-]*\d{3}[\s\-]*\d{2}[\s\-]*\d{2}"
    r")(?!\d)"
)
RE_PHONE_AREA_CONTEXT = re.compile(
    r"\(\s*[0-9OoОо]{3}\s*\)"
    r"(?=[^\n]{0,140}(?:\+7|8\b|[0-9OoОо]{3}\s*[-\s]*[0-9OoОо]{2}\s*[-\s]*[0-9OoОо]{2}|телефон|phone))",
    re.IGNORECASE,
)
RE_INN = re.compile(r"(?<!\d)\d{10}(?!\d)|(?<!\d)\d{12}(?!\d)")
RE_PASSPORT = re.compile(r"(?i)\b(?:паспорт|passport)\b[^0-9]{0,24}(\d{2})\s*(\d{2})\s*(\d{6})")
RE_MONEY = re.compile(
    r"(?<!\w)(\d[\d\s]{0,15})(?:[.,]\d{1,2})?\s*(?:руб\.?|₽|RUB|usd|USD|eur|EUR)(?=\s|$|[.,;:)\]])",
    re.IGNORECASE,
)
RE_ACCOUNT = re.compile(r"(?<!\d)(?:\d[\s\-]?){20}(?!\d)")
RE_ORG = re.compile(
    r"\b(?:ООО|ОАО|ПАО|АО|ЗАО|ИП)\s+[\"«]?[A-Za-zА-Яа-яЁё0-9][A-Za-zА-Яа-яЁё0-9 \t\-]{1,80}?[\"»]?(?=$|[.,;:\n)])",
    re.IGNORECASE,
)

REGEX_DETECTORS = {
    "ORG": RE_ORG,
    "EMAIL": RE_EMAIL,
    "PHONE": RE_PHONE,
    "INN": RE_INN,
    "PASSPORT": RE_PASSPORT,
    "MONEY": RE_MONEY,
    "ACCOUNT": RE_ACCOUNT,
}

LABEL_PRIORITY = {
    "ACCOUNT": 9,
    "PASSPORT": 8,
    "INN": 7,
    "EMAIL": 6,
    "PHONE": 5,
    "MONEY": 4,
    "ORG": 3,
    "PER": 2,
    "LOC": 1,
    "CUSTOM": 1,
}

LABEL_TITLES = {
    "PER": "ФИО",
    "ORG": "Организация",
    "LOC": "Локация",
    "EMAIL": "Email",
    "PHONE": "Телефон",
    "MONEY": "Сумма",
    "PASSPORT": "Паспорт",
    "ACCOUNT": "Счет",
    "INN": "ИНН",
    "CUSTOM": "Словарь",
    "MANUAL": "Вручную",
}

LABEL_PREVIEW_COLORS = {
    "PER": ("#fff1f2", "#be123c"),
    "ORG": ("#eff6ff", "#1d4ed8"),
    "LOC": ("#ecfdf5", "#047857"),
    "EMAIL": ("#f5f3ff", "#6d28d9"),
    "PHONE": ("#fff7ed", "#c2410c"),
    "MONEY": ("#fefce8", "#a16207"),
    "PASSPORT": ("#f3e8ff", "#7e22ce"),
    "ACCOUNT": ("#ecfeff", "#0f766e"),
    "INN": ("#f0fdf4", "#15803d"),
    "CUSTOM": ("#f8fafc", "#475569"),
    "MANUAL": ("#e0f2fe", "#0369a1"),
}


@dataclass(frozen=True)
class Span:
    start: int
    end: int
    label: str
    text: str
    page: Optional[int] = None
    method: str = "TextLayer"


class UnsupportedFileError(ValueError):
    pass


def get_extension(filename: str) -> str:
    return filename.rsplit(".", 1)[-1].lower() if "." in filename else ""


def safe_name(filename: str) -> str:
    return os.path.basename(filename).strip() or "file"


def normalize_for_match(value: str) -> str:
    value = _normalize_ocr_text_for_regex(value or "")
    return re.sub(r"\s+", " ", value).strip().lower()


def _replace_ocr_space_with_dot(match: re.Match[str]) -> str:
    left = match.group(1)
    spaces = match.group(2)
    right = match.group(3)
    # Keep string length unchanged for start/end mapping.
    return f"{left}.{spaces[1:]}{right}" if spaces else f"{left}.{right}"


def _normalize_ocr_text_for_regex(text: str) -> str:
    if not text:
        return text

    normalized = text

    # Typical OCR confusion in numeric fragments: O/o/О/о -> 0
    normalized = re.sub(r"(?<=\d)[OoОо](?=\d)", "0", normalized)
    normalized = re.sub(r"(?<=\d)[OoОо](?=[\s\-\)])", "0", normalized)
    normalized = re.sub(r"(?<=[\s\-\(])[OoОо](?=\d)", "0", normalized)

    # Email local-part: "smirnov a@" -> "smirnov.a@"
    normalized = re.sub(
        r"([A-Za-z0-9._%+\-])(\s+)([A-Za-z0-9._%+\-])(?=\s*@)",
        _replace_ocr_space_with_dot,
        normalized,
    )

    # Email domain: "@gazprom ru" -> "@gazprom.ru"
    normalized = re.sub(
        r"(@[A-Za-z0-9.\-]+)(\s+)([A-Za-z]{2,})(?=[\s:;,.!?)]|$)",
        _replace_ocr_space_with_dot,
        normalized,
    )

    return normalized


def parse_custom_words(raw: str) -> List[str]:
    if not raw:
        return []
    return [line.strip() for line in raw.splitlines() if line.strip()]


def expand_archives(files: Sequence[Tuple[str, bytes]]) -> List[Tuple[str, bytes]]:
    expanded: List[Tuple[str, bytes]] = []
    for filename, data in files:
        name = safe_name(filename)
        ext = get_extension(name)
        if ext != "zip":
            expanded.append((name, data))
            continue

        with zipfile.ZipFile(io.BytesIO(data)) as zf:
            for member in zf.namelist():
                if member.endswith("/"):
                    continue
                member_name = safe_name(member)
                member_ext = get_extension(member_name)
                if member_ext in SUPPORTED_EXTENSIONS and member_ext != "zip":
                    expanded.append((member_name, zf.read(member)))
    return expanded


@lru_cache(maxsize=1)
def load_natasha() -> Tuple[Segmenter, NewsNerTagger]:
    embedding = NewsEmbedding()
    return Segmenter(), NewsNerTagger(embedding)


@lru_cache(maxsize=1)
def load_ocr_reader():
    if not OCR_AVAILABLE:
        return None
    use_gpu = os.getenv("OCR_GPU", "0").strip().lower() in {"1", "true", "yes", "on"}
    return easyocr.Reader(["ru", "en"], gpu=use_gpu)


def qwen_is_configured() -> bool:
    return bool(os.getenv("QWEN_API_BASE") and os.getenv("QWEN_MODEL"))


def resolve_engine(engine_preference: str, text: str) -> str:
    preferred = (engine_preference or "auto").strip().lower()
    if preferred in {"natasha", "regex"}:
        return preferred
    if preferred == "qwen":
        if qwen_is_configured():
            return "qwen"
        preferred = "auto"

    ru_letters = len(re.findall(r"[А-Яа-яЁё]", text))
    en_letters = len(re.findall(r"[A-Za-z]", text))

    if ru_letters >= en_letters:
        return "natasha"
    if qwen_is_configured():
        return "qwen"
    return "regex"


def _extract_json_from_text(raw: str) -> Optional[object]:
    content = (raw or "").strip()
    if not content:
        return None

    if content.startswith("```"):
        content = re.sub(r"^```[a-zA-Z]*\n", "", content)
        content = re.sub(r"\n```$", "", content)

    try:
        return json.loads(content)
    except json.JSONDecodeError:
        match = re.search(r"\[[\s\S]*\]", content)
        if match:
            try:
                return json.loads(match.group(0))
            except json.JSONDecodeError:
                return None
    return None


def qwen_ner(text: str, categories: Set[str], page: Optional[int], method: str) -> List[Span]:
    if not qwen_is_configured() or not text.strip():
        return []

    endpoint = os.getenv("QWEN_API_BASE", "").rstrip("/") + "/chat/completions"
    model = os.getenv("QWEN_MODEL", "")
    api_key = os.getenv("QWEN_API_KEY", "")

    prompt = (
        "Извлеки сущности из текста. Верни ТОЛЬКО JSON-массив объектов. "
        "Каждый объект: {\"label\": \"...\", \"text\": \"...\"}. "
        "Допустимые label: PER, ORG, LOC, EMAIL, PHONE, INN, PASSPORT, MONEY, ACCOUNT."
    )

    headers = {"Content-Type": "application/json"}
    if api_key:
        headers["Authorization"] = f"Bearer {api_key}"

    payload = {
        "model": model,
        "temperature": 0,
        "messages": [
            {"role": "system", "content": "Ты NER-модуль. Отвечай только JSON."},
            {"role": "user", "content": f"{prompt}\n\nТекст:\n{text}"},
        ],
    }

    try:
        response = requests.post(endpoint, headers=headers, json=payload, timeout=45)
        response.raise_for_status()
        content = response.json()["choices"][0]["message"]["content"]
        parsed = _extract_json_from_text(content)
    except Exception:
        return []

    if not isinstance(parsed, list):
        return []

    spans: List[Span] = []
    for item in parsed:
        if not isinstance(item, dict):
            continue
        label = str(item.get("label", "")).upper().strip()
        fragment = str(item.get("text", "")).strip()
        if not label or not fragment:
            continue
        if label not in categories:
            continue

        for match in re.finditer(re.escape(fragment), text, flags=re.IGNORECASE):
            spans.append(
                Span(
                    start=match.start(),
                    end=match.end(),
                    label=label,
                    text=text[match.start() : match.end()],
                    page=page,
                    method=method,
                )
            )
            break

    return spans


def find_spans(
    text: str,
    categories: Set[str],
    custom_words: Sequence[str],
    engine_preference: str,
    page: Optional[int] = None,
    method: str = "TextLayer",
) -> Tuple[List[Span], str]:
    if not text.strip():
        return [], resolve_engine(engine_preference, text)

    engine_used = resolve_engine(engine_preference, text)
    spans: List[Span] = []

    if engine_used == "natasha":
        segmenter, ner_tagger = load_natasha()
        doc = Doc(text)
        doc.segment(segmenter)
        doc.tag_ner(ner_tagger)
        for item in doc.spans:
            label = str(item.type).upper()
            if label in categories:
                spans.append(
                    Span(
                        start=item.start,
                        end=item.stop,
                        label=label,
                        text=text[item.start : item.stop],
                        page=page,
                        method=method,
                    )
                )
    elif engine_used == "qwen":
        spans.extend(qwen_ner(text, categories, page, method))

    regex_text = _normalize_ocr_text_for_regex(text) if method == "OCR" else text

    for label, regex in REGEX_DETECTORS.items():
        if label not in categories:
            continue
        for match in regex.finditer(regex_text):
            spans.append(
                Span(
                    start=match.start(),
                    end=match.end(),
                    label=label,
                    text=text[match.start() : match.end()],
                    page=page,
                    method=method,
                )
            )

    if method == "OCR" and "PHONE" in categories:
        for match in RE_PHONE_AREA_CONTEXT.finditer(regex_text):
            spans.append(
                Span(
                    start=match.start(),
                    end=match.end(),
                    label="PHONE",
                    text=text[match.start() : match.end()],
                    page=page,
                    method=method,
                )
            )

    if "CUSTOM" in categories and custom_words:
        for term in custom_words:
            for match in re.finditer(re.escape(term), text, flags=re.IGNORECASE):
                spans.append(
                    Span(
                        start=match.start(),
                        end=match.end(),
                        label="CUSTOM",
                        text=text[match.start() : match.end()],
                        page=page,
                        method=method,
                    )
                )

    spans = sorted(
        spans,
        key=lambda item: (
            -(item.end - item.start),
            -LABEL_PRIORITY.get(item.label, 0),
            item.start,
        ),
    )
    picked: List[Span] = []
    for span in spans:
        overlap = any(not (span.end <= current.start or current.end <= span.start) for current in picked)
        if not overlap:
            picked.append(span)

    normalized_picked: List[Span] = []
    for span in picked:
        if span.method == "OCR" and span.label == "ORG" and "\n" in span.text:
            trimmed = span.text.split("\n", 1)[0].rstrip(" ,;:")
            if trimmed:
                normalized_picked.append(
                    Span(
                        start=span.start,
                        end=span.start + len(trimmed),
                        label=span.label,
                        text=text[span.start : span.start + len(trimmed)],
                        page=span.page,
                        method=span.method,
                    )
                )
                continue
        normalized_picked.append(span)

    return sorted(normalized_picked, key=lambda item: item.start), engine_used


def _bbox_to_xy(bbox: List[List[float]]) -> Tuple[float, float, float, float, float, float]:
    xs = [point[0] for point in bbox]
    ys = [point[1] for point in bbox]
    x1, x2 = min(xs), max(xs)
    y1, y2 = min(ys), max(ys)
    return x1, y1, x2, y2, (x1 + x2) / 2, (y1 + y2) / 2


def _group_tokens_by_lines(tokens: List[Dict[str, float]], image_height: int) -> List[Dict[str, object]]:
    tokens = sorted(tokens, key=lambda token: (token["cy"], token["x1"]))
    lines: List[Dict[str, object]] = []
    y_tolerance = max(12, int(image_height * 0.015))

    for token in tokens:
        placed = False
        for line in lines:
            if abs(token["cy"] - line["cy"]) <= y_tolerance:
                line["tokens"].append(token)
                line["cy"] = (line["cy"] * 0.8) + (token["cy"] * 0.2)
                placed = True
                break
        if not placed:
            lines.append({"cy": token["cy"], "tokens": [token]})

    return lines


def _ocr_lines_from_image(image: np.ndarray) -> List[Dict[str, object]]:
    reader = load_ocr_reader()
    if reader is None:
        return []

    raw = reader.readtext(image, detail=1, paragraph=False)
    tokens: List[Dict[str, float]] = []
    for bbox, text, _ in raw:
        x1, y1, x2, y2, cx, cy = _bbox_to_xy(bbox)
        tokens.append(
            {
                "bbox": bbox,
                "text": text,
                "x1": x1,
                "y1": y1,
                "x2": x2,
                "y2": y2,
                "cx": cx,
                "cy": cy,
            }
        )

    lines: List[Dict[str, object]] = []
    for grouped in _group_tokens_by_lines(tokens, image.shape[0]):
        parts: List[str] = []
        position = 0
        line_tokens = sorted(grouped["tokens"], key=lambda token: (float(token["x1"]), float(token["cx"])))
        for token in line_tokens:
            token["s"] = position
            parts.append(token["text"])
            position += len(token["text"]) + 1
            token["e"] = position - 1

        lines.append({"text": " ".join(parts), "tokens": line_tokens})

    return lines


def _flatten_ocr_lines(lines: Sequence[Dict[str, object]]) -> Tuple[str, List[Dict[str, object]]]:
    full_text_parts: List[str] = []
    flat_tokens: List[Dict[str, object]] = []
    offset = 0

    for line in lines:
        line_text = str(line["text"])
        full_text_parts.append(line_text)

        for token in line["tokens"]:
            token_copy = dict(token)
            token_copy["gs"] = offset + int(token["s"])
            token_copy["ge"] = offset + int(token["e"])
            flat_tokens.append(token_copy)

        offset += len(line_text) + 1

    return "\n".join(full_text_parts), flat_tokens


def _token_rects_for_hit(
    flat_tokens: Sequence[Dict[str, object]],
    start: int,
    end: int,
    x_scale: float = 1.0,
    y_scale: float = 1.0,
    pad: int = 4,
) -> List[Tuple[float, float, float, float]]:
    rects: List[Tuple[float, float, float, float]] = []

    for token in flat_tokens:
        token_start = int(token["gs"])
        token_end = int(token["ge"])
        if token_start >= end or token_end <= start:
            continue

        x1 = float(token["x1"]) * x_scale
        y1 = float(token["y1"]) * y_scale
        x2 = float(token["x2"]) * x_scale
        y2 = float(token["y2"]) * y_scale

        token_text = str(token.get("text", ""))
        token_len = max(token_end - token_start, len(token_text), 1)
        local_start = max(0, start - token_start)
        local_end = min(token_len, end - token_start)
        if local_end <= local_start:
            continue

        if local_start == 0 and local_end >= token_len:
            rect_x1, rect_x2 = x1, x2
        else:
            width = max(1.0, x2 - x1)
            char_width = width / float(token_len)
            rect_x1 = x1 + (char_width * local_start)
            rect_x2 = x1 + (char_width * local_end)

            min_width = max(8.0 * x_scale, char_width * 2.0)
            if rect_x2 - rect_x1 < min_width:
                center = (rect_x1 + rect_x2) / 2.0
                rect_x1 = center - (min_width / 2.0)
                rect_x2 = center + (min_width / 2.0)

        rects.append((rect_x1 - pad, y1 - pad, rect_x2 + pad, y2 + pad))

    return rects


def _extract_pdf_text(page: fitz.Page, use_ocr: bool) -> Tuple[str, str]:
    text = page.get_text("text")
    if text.strip() or not use_ocr:
        return text, "TextLayer"

    reader = load_ocr_reader()
    if reader is None:
        return "", "TextLayer"

    pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
    img = np.frombuffer(pix.samples, dtype=np.uint8).reshape((pix.height, pix.width, pix.n))
    if pix.n == 4:
        img = img[:, :, :3]

    lines = _ocr_lines_from_image(img)
    text = "\n".join(line["text"] for line in lines)
    return text, "OCR"


def apply_text_redaction(text: str, hits: Sequence[Dict[str, object]], style: str = "black") -> str:
    replacement_by_style = {
        "black": "██████",
        "tag": None,
    }

    redacted = text
    for hit in sorted(hits, key=lambda item: int(item["start"]), reverse=True):
        start = int(hit["start"])
        end = int(hit["end"])
        if start < 0 or end > len(redacted) or start >= end:
            continue

        replacement = replacement_by_style.get(style)
        if replacement is None:
            replacement = f"[{hit['label']}]"

        redacted = redacted[:start] + replacement + redacted[end:]

    return redacted


def _summarize_hits(hits: Sequence[Dict[str, object]]) -> Dict[str, int]:
    summary: Dict[str, int] = defaultdict(int)
    for hit in hits:
        summary[str(hit["label"])] += 1
    return dict(sorted(summary.items(), key=lambda item: item[0]))


def _build_verdict(hits_count: int) -> str:
    if hits_count > 0:
        return "Нельзя передавать без обезличивания"
    return "Можно передавать"


def _spans_to_hits(spans: Sequence[Span]) -> List[Dict[str, object]]:
    hits: List[Dict[str, object]] = []
    for index, span in enumerate(spans, start=1):
        hits.append(
            {
                "id": f"h{index}",
                "label": span.label,
                "text": span.text,
                "page": span.page,
                "start": span.start,
                "end": span.end,
                "method": span.method,
            }
        )
    return hits


def analyze_file(
    file_id: str,
    filename: str,
    data: bytes,
    categories: Set[str],
    custom_words: Sequence[str],
    use_ocr: bool,
    engine_preference: str,
) -> Tuple[Dict[str, object], Dict[str, object]]:
    extension = get_extension(filename)
    if extension not in SUPPORTED_EXTENSIONS or extension == "zip":
        raise UnsupportedFileError(f"Неподдерживаемый формат: {filename}")

    if extension == "pdf":
        analysis, state = _analyze_pdf(file_id, filename, data, categories, custom_words, use_ocr, engine_preference)
    elif extension == "docx":
        analysis, state = _analyze_docx(file_id, filename, data, categories, custom_words, engine_preference)
    elif extension in {"txt", "md"}:
        analysis, state = _analyze_text(file_id, filename, data, categories, custom_words, engine_preference)
    elif extension in {"png", "jpg", "jpeg"}:
        analysis, state = _analyze_image(file_id, filename, data, categories, custom_words, engine_preference)
    else:
        raise UnsupportedFileError(f"Неподдерживаемый формат: {filename}")

    return analysis, state


def _build_preview_for_pdf(page_texts: Sequence[str], page_hits: Dict[int, List[Dict[str, object]]]) -> str:
    blocks: List[str] = []
    for page_index, text in enumerate(page_texts, start=1):
        if not text.strip():
            continue
        hits = page_hits.get(page_index, [])
        redacted = apply_text_redaction(text, hits, style="tag")
        blocks.append(f"### Страница {page_index}\n{redacted}")
        if len("\n\n".join(blocks)) > 6000:
            break

    return "\n\n".join(blocks)[:6000]


def _replace_newlines_for_html(value: str) -> str:
    return value.replace("\n", "<br>")


def _highlight_text_html(text: str, hits: Sequence[Dict[str, object]], max_chars: int = 3000) -> str:
    if not text.strip():
        return "<div class='preview-empty'>Пустой фрагмент</div>"

    text = text[:max_chars]
    safe_parts: List[str] = []
    cursor = 0

    for hit in sorted(hits, key=lambda item: int(item["start"])):
        start = max(0, min(int(hit["start"]), len(text)))
        end = max(start, min(int(hit["end"]), len(text)))
        if start >= len(text) or end <= cursor:
            continue

        if start > cursor:
            safe_parts.append(html.escape(text[cursor:start]))

        bg_color, fg_color = LABEL_PREVIEW_COLORS.get(str(hit["label"]), ("#f8fafc", "#334155"))
        safe_parts.append(
            (
                f"<mark class='preview-hit' "
                f"style='background:{bg_color}; color:{fg_color}; border-color:{fg_color}22;'>"
                f"<span class='preview-hit__tag'>{html.escape(LABEL_TITLES.get(str(hit['label']), str(hit['label'])))}"
                f"</span>{html.escape(text[start:end])}</mark>"
            )
        )
        cursor = end

    if cursor < len(text):
        safe_parts.append(html.escape(text[cursor:]))

    body = _replace_newlines_for_html("".join(safe_parts))
    return f"<div class='preview-surface'>{body}</div>"


def _build_preview_html_for_pdf(page_texts: Sequence[str], page_hits: Dict[int, List[Dict[str, object]]]) -> str:
    blocks: List[str] = []
    remaining_chars = 5000

    for page_index, text in enumerate(page_texts, start=1):
        if not text.strip():
            continue

        chunk = text[:remaining_chars]
        blocks.append(
            "<section class='preview-page'>"
            f"<div class='preview-page__title'>Страница {page_index}</div>"
            f"{_highlight_text_html(chunk, page_hits.get(page_index, []), max_chars=remaining_chars)}"
            "</section>"
        )
        remaining_chars -= len(chunk)
        if remaining_chars <= 0:
            break

    if not blocks:
        return "<div class='preview-empty'>Предпросмотр недоступен</div>"

    return "".join(blocks)


def _build_preview_pages_for_pdf(
    page_texts: Sequence[str],
    page_hits: Dict[int, List[Dict[str, object]]],
    max_chars_per_page: int = 2200,
) -> List[Dict[str, object]]:
    pages: List[Dict[str, object]] = []

    for page_index, text in enumerate(page_texts, start=1):
        if not text.strip():
            html_block = "<div class='preview-empty'>Пустая страница</div>"
        else:
            html_block = _highlight_text_html(text, page_hits.get(page_index, []), max_chars=max_chars_per_page)

        pages.append(
            {
                "page": page_index,
                "html": (
                    "<section class='preview-page'>"
                    f"<div class='preview-page__title'>Страница {page_index}</div>"
                    f"{html_block}"
                    "</section>"
                ),
            }
        )

    return pages


def _analyze_pdf(
    file_id: str,
    filename: str,
    data: bytes,
    categories: Set[str],
    custom_words: Sequence[str],
    use_ocr: bool,
    engine_preference: str,
) -> Tuple[Dict[str, object], Dict[str, object]]:
    doc = fitz.open(stream=data, filetype="pdf")
    spans: List[Span] = []
    page_texts: List[str] = []
    engines_used: Set[str] = set()

    for page_index in range(doc.page_count):
        page = doc.load_page(page_index)
        text, method = _extract_pdf_text(page, use_ocr)
        page_texts.append(text)
        page_spans, engine_used = find_spans(
            text,
            categories,
            custom_words,
            engine_preference,
            page=page_index + 1,
            method=method,
        )
        spans.extend(page_spans)
        engines_used.add(engine_used)

    doc.close()

    hits = _spans_to_hits(spans)
    hits_by_page: Dict[int, List[Dict[str, object]]] = defaultdict(list)
    for hit in hits:
        if hit["page"] is not None:
            hits_by_page[int(hit["page"])].append(hit)

    preview = _build_preview_for_pdf(page_texts, hits_by_page)
    preview_html = _build_preview_html_for_pdf(page_texts, hits_by_page)

    analysis = {
        "file_id": file_id,
        "filename": filename,
        "extension": "pdf",
        "engine_used": ", ".join(sorted(engines_used)) if engines_used else "regex",
        "hits": hits,
        "summary": _summarize_hits(hits),
        "verdict": _build_verdict(len(hits)),
        "preview": preview,
        "preview_html": preview_html,
        "page_count": len(page_texts),
        "preview_pages": _build_preview_pages_for_pdf(page_texts, hits_by_page),
    }

    state = {
        "file_id": file_id,
        "filename": filename,
        "extension": "pdf",
        "bytes": data,
        "hits": hits,
        "page_texts": page_texts,
        "engine_preference": engine_preference,
    }
    return analysis, state


def _extract_docx_text(data: bytes) -> str:
    candidate_names = (
        "word/document.xml",
        "word/footnotes.xml",
        "word/endnotes.xml",
        "word/header1.xml",
        "word/header2.xml",
        "word/header3.xml",
        "word/footer1.xml",
        "word/footer2.xml",
        "word/footer3.xml",
    )
    paragraphs: List[str] = []

    with zipfile.ZipFile(io.BytesIO(data)) as archive:
        for name in candidate_names:
            if name not in archive.namelist():
                continue
            try:
                root = ET.fromstring(archive.read(name))
            except ET.ParseError:
                continue

            for paragraph in root.iter():
                if not paragraph.tag.endswith("}p"):
                    continue
                parts = [node.text for node in paragraph.iter() if node.tag.endswith("}t") and node.text]
                if parts:
                    paragraphs.append("".join(parts).strip())

    return "\n".join(item for item in paragraphs if item)


def _analyze_docx(
    file_id: str,
    filename: str,
    data: bytes,
    categories: Set[str],
    custom_words: Sequence[str],
    engine_preference: str,
) -> Tuple[Dict[str, object], Dict[str, object]]:
    try:
        doc = Document(io.BytesIO(data))
        full_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    except Exception:
        full_text = _extract_docx_text(data)
        if not full_text.strip():
            raise UnsupportedFileError(f"Не удалось прочитать содержимое DOCX: {filename}")

    spans, engine_used = find_spans(full_text, categories, custom_words, engine_preference, page=1, method="Text")
    hits = _spans_to_hits(spans)

    analysis = {
        "file_id": file_id,
        "filename": filename,
        "extension": "docx",
        "engine_used": engine_used,
        "hits": hits,
        "summary": _summarize_hits(hits),
        "verdict": _build_verdict(len(hits)),
        "preview": apply_text_redaction(full_text, hits, style="tag")[:6000],
        "preview_html": _highlight_text_html(full_text, hits, max_chars=5000),
    }

    state = {
        "file_id": file_id,
        "filename": filename,
        "extension": "docx",
        "bytes": data,
        "hits": hits,
        "text": full_text,
        "engine_preference": engine_preference,
    }
    return analysis, state


def _analyze_text(
    file_id: str,
    filename: str,
    data: bytes,
    categories: Set[str],
    custom_words: Sequence[str],
    engine_preference: str,
) -> Tuple[Dict[str, object], Dict[str, object]]:
    full_text = data.decode("utf-8", errors="replace")

    spans, engine_used = find_spans(full_text, categories, custom_words, engine_preference, page=1, method="Text")
    hits = _spans_to_hits(spans)

    analysis = {
        "file_id": file_id,
        "filename": filename,
        "extension": get_extension(filename),
        "engine_used": engine_used,
        "hits": hits,
        "summary": _summarize_hits(hits),
        "verdict": _build_verdict(len(hits)),
        "preview": apply_text_redaction(full_text, hits, style="tag")[:6000],
        "preview_html": _highlight_text_html(full_text, hits, max_chars=5000),
    }

    state = {
        "file_id": file_id,
        "filename": filename,
        "extension": get_extension(filename),
        "bytes": data,
        "hits": hits,
        "text": full_text,
        "engine_preference": engine_preference,
    }
    return analysis, state


def _analyze_image(
    file_id: str,
    filename: str,
    data: bytes,
    categories: Set[str],
    custom_words: Sequence[str],
    engine_preference: str,
) -> Tuple[Dict[str, object], Dict[str, object]]:
    array = np.frombuffer(data, np.uint8)
    image = cv2.imdecode(array, cv2.IMREAD_COLOR)
    if image is None:
        raise UnsupportedFileError(f"Не удалось прочитать изображение: {filename}")

    lines = _ocr_lines_from_image(image)
    full_text, _ = _flatten_ocr_lines(lines)
    spans, engine_used = find_spans(
        full_text,
        categories,
        custom_words,
        engine_preference,
        page=1,
        method="OCR",
    )
    hits = _spans_to_hits(spans)

    analysis = {
        "file_id": file_id,
        "filename": filename,
        "extension": get_extension(filename),
        "engine_used": engine_used,
        "hits": hits,
        "summary": _summarize_hits(hits),
        "verdict": _build_verdict(len(hits)),
        "preview": apply_text_redaction(full_text, hits, style="tag")[:6000],
        "preview_html": _highlight_text_html(full_text, hits, max_chars=5000),
    }

    state = {
        "file_id": file_id,
        "filename": filename,
        "extension": get_extension(filename),
        "bytes": data,
        "hits": hits,
        "text": full_text,
        "engine_preference": engine_preference,
    }
    return analysis, state


def _build_docx(text: str, title: Optional[str] = None) -> bytes:
    document = Document()
    if title:
        document.add_heading(title, level=1)

    for line in text.splitlines() or [text]:
        document.add_paragraph(line)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _markdown_escape(text: str) -> str:
    return text.replace("|", "\\|")


def _create_markdown_report(report: Dict[str, List[Dict[str, object]]]) -> str:
    lines: List[str] = ["# Отчёт по обезличиванию", ""]
    for filename, items in report.items():
        lines.append(f"## {filename}")
        if not items:
            lines.append("- Совпадений для вычеркивания не выбрано")
            lines.append("")
            continue

        lines.append("| Страница | Тип | Текст | Метод |")
        lines.append("|---:|---|---|---|")
        for item in items:
            page = item.get("page") if item.get("page") is not None else "-"
            lines.append(
                f"| {page} | {item['label']} | {_markdown_escape(str(item['text']))} | {item.get('method', '-')} |"
            )
        lines.append("")

    return "\n".join(lines)


def _dedupe_hits(hits: Sequence[Dict[str, object]]) -> List[Dict[str, object]]:
    seen: Set[Tuple[object, ...]] = set()
    unique: List[Dict[str, object]] = []

    for hit in hits:
        key = (
            hit.get("page"),
            hit.get("start"),
            hit.get("end"),
            hit.get("label"),
            hit.get("text"),
            hit.get("method"),
        )
        if key in seen:
            continue
        seen.add(key)
        unique.append(hit)

    return unique


def _find_manual_hits(file_state: Dict[str, object], manual_terms: Sequence[str]) -> List[Dict[str, object]]:
    clean_terms = [term.strip() for term in manual_terms if term and term.strip()]
    if not clean_terms:
        return []

    hits: List[Dict[str, object]] = []
    extension = str(file_state.get("extension", ""))

    if extension == "pdf":
        counter = 1
        for page_index, page_text in enumerate(file_state.get("page_texts", []), start=1):
            for term in clean_terms:
                for match in re.finditer(re.escape(term), str(page_text), flags=re.IGNORECASE):
                    hits.append(
                        {
                            "id": f"m{counter}",
                            "label": "MANUAL",
                            "text": str(page_text)[match.start() : match.end()],
                            "page": page_index,
                            "start": match.start(),
                            "end": match.end(),
                            "method": "Manual",
                        }
                    )
                    counter += 1
        return _dedupe_hits(hits)

    source_text = str(file_state.get("text", ""))
    counter = 1
    for term in clean_terms:
        for match in re.finditer(re.escape(term), source_text, flags=re.IGNORECASE):
            hits.append(
                {
                    "id": f"m{counter}",
                    "label": "MANUAL",
                    "text": source_text[match.start() : match.end()],
                    "page": 1,
                    "start": match.start(),
                    "end": match.end(),
                    "method": "Manual",
                }
            )
            counter += 1

    return _dedupe_hits(hits)


def _build_output_name(filename: str, suffix: str, extension: str) -> str:
    stem = os.path.splitext(safe_name(filename))[0]
    return f"{stem}{suffix}.{extension}"


def _add_unique_output(
    output_files: List[Tuple[str, bytes]],
    used_names: Set[str],
    filename: str,
    payload: bytes,
) -> None:
    if filename not in used_names:
        used_names.add(filename)
        output_files.append((filename, payload))
        return

    stem, ext = os.path.splitext(filename)
    counter = 2
    while True:
        candidate = f"{stem}_{counter}{ext}"
        if candidate not in used_names:
            used_names.add(candidate)
            output_files.append((candidate, payload))
            return
        counter += 1


def _redact_image(
    image_bytes: bytes,
    selected_hits: Sequence[Dict[str, object]],
    categories: Set[str],
    custom_words: Sequence[str],
    engine_preference: str,
) -> bytes:
    array = np.frombuffer(image_bytes, np.uint8)
    image = cv2.imdecode(array, cv2.IMREAD_COLOR)
    if image is None:
        return image_bytes

    lines = _ocr_lines_from_image(image)
    _, flat_tokens = _flatten_ocr_lines(lines)

    for hit in selected_hits:
        start = int(hit.get("start", -1))
        end = int(hit.get("end", -1))
        if start < 0 or end <= start:
            continue

        hit_rects = _token_rects_for_hit(flat_tokens, start, end, pad=4)
        if not hit_rects:
            continue

        for x1, y1, x2, y2 in hit_rects:
            cv2.rectangle(
                image,
                (max(0, math.floor(x1)), max(0, math.floor(y1))),
                (min(image.shape[1], math.ceil(x2)), min(image.shape[0], math.ceil(y2))),
                (0, 0, 0),
                -1,
            )

    success, buffer = cv2.imencode(".png", image)
    return buffer.tobytes() if success else image_bytes


def _redact_pdf(
    pdf_bytes: bytes,
    selected_hits: Sequence[Dict[str, object]],
    categories: Set[str],
    custom_words: Sequence[str],
    use_ocr: bool,
    engine_preference: str,
) -> bytes:
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")

    selected_by_page: Dict[int, List[Dict[str, object]]] = defaultdict(list)
    for hit in selected_hits:
        page = int(hit.get("page") or 1)
        selected_by_page[page].append(hit)

    for page_index in range(doc.page_count):
        page_number = page_index + 1
        page = doc.load_page(page_index)
        page_hits = selected_by_page.get(page_number, [])
        if not page_hits:
            continue

        text = page.get_text("text")
        if text.strip():
            for hit in page_hits:
                for rect in page.search_for(str(hit["text"])):
                    page.add_redact_annot(rect, fill=(0, 0, 0))
            page.apply_redactions()
            continue

        if not use_ocr:
            continue

        reader = load_ocr_reader()
        if reader is None:
            continue

        pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
        img = np.frombuffer(pix.samples, dtype=np.uint8).reshape((pix.height, pix.width, pix.n))
        if pix.n == 4:
            img = img[:, :, :3]

        lines = _ocr_lines_from_image(img)
        scale_x = page.rect.width / float(pix.width)
        scale_y = page.rect.height / float(pix.height)
        _, flat_tokens = _flatten_ocr_lines(lines)

        for hit in page_hits:
            start = int(hit.get("start", -1))
            end = int(hit.get("end", -1))
            if start < 0 or end <= start:
                continue

            hit_rects = _token_rects_for_hit(flat_tokens, start, end, x_scale=scale_x, y_scale=y_scale, pad=4)
            if not hit_rects:
                continue

            for x1, y1, x2, y2 in hit_rects:
                rect = fitz.Rect(x1, y1, x2, y2)
                page.add_redact_annot(rect, fill=(0, 0, 0))

        page.apply_redactions()

    result = doc.write()
    doc.close()
    return result


def build_redacted_zip(
    analysis_state: Dict[str, object],
    selected_hit_ids_by_file: Dict[str, List[str]],
    manual_terms_by_file: Optional[Dict[str, List[str]]] = None,
    redaction_style: str = "black",
    include_original: bool = True,
    include_markdown: bool = True,
    include_docx: bool = True,
) -> Tuple[bytes, Dict[str, List[Dict[str, object]]]]:
    categories = set(analysis_state["categories"])
    custom_words = list(analysis_state["custom_words"])
    use_ocr = bool(analysis_state["use_ocr"])
    engine_preference = str(analysis_state["engine"])

    report: Dict[str, List[Dict[str, object]]] = {}
    output_files: List[Tuple[str, bytes]] = []
    used_output_names: Set[str] = set()

    for file_state in analysis_state["files"]:
        file_id = str(file_state["file_id"])
        filename = str(file_state["filename"])
        extension = str(file_state["extension"])

        selected_ids = set(selected_hit_ids_by_file.get(file_id, []))
        selected_hits = [hit for hit in file_state["hits"] if hit["id"] in selected_ids]
        manual_hits = _find_manual_hits(file_state, (manual_terms_by_file or {}).get(file_id, []))
        selected_hits = _dedupe_hits(selected_hits + manual_hits)

        report[filename] = [
            {
                "page": hit.get("page"),
                "label": hit["label"],
                "text": hit["text"],
                "method": hit.get("method"),
            }
            for hit in selected_hits
        ]

        if extension == "pdf":
            redacted_pdf = _redact_pdf(
                file_state["bytes"],
                selected_hits,
                categories,
                custom_words,
                use_ocr,
                engine_preference,
            )

            if include_original:
                _add_unique_output(
                    output_files,
                    used_output_names,
                    _build_output_name(filename, "_redacted", "pdf"),
                    redacted_pdf,
                )

            page_texts = file_state.get("page_texts", [])
            page_groups: Dict[int, List[Dict[str, object]]] = defaultdict(list)
            for hit in selected_hits:
                page_groups[int(hit.get("page") or 1)].append(hit)

            redacted_pages: List[str] = []
            for page_index, page_text in enumerate(page_texts, start=1):
                page_redacted = apply_text_redaction(page_text, page_groups.get(page_index, []), style=redaction_style)
                redacted_pages.append(f"## Страница {page_index}\n\n{page_redacted}")

            redacted_text = "\n\n".join(redacted_pages)
            if include_markdown:
                _add_unique_output(
                    output_files,
                    used_output_names,
                    _build_output_name(filename, "_redacted", "md"),
                    redacted_text.encode("utf-8"),
                )
            if include_docx:
                _add_unique_output(
                    output_files,
                    used_output_names,
                    _build_output_name(filename, "_redacted", "docx"),
                    _build_docx(redacted_text),
                )
            continue

        if extension in {"docx", "txt", "md"}:
            source_text = str(file_state.get("text", ""))
            redacted_text = apply_text_redaction(source_text, selected_hits, style=redaction_style)
            original_written = False

            if include_original:
                if extension == "docx":
                    _add_unique_output(
                        output_files,
                        used_output_names,
                        _build_output_name(filename, "_redacted", "docx"),
                        _build_docx(redacted_text),
                    )
                else:
                    _add_unique_output(
                        output_files,
                        used_output_names,
                        _build_output_name(filename, "_redacted", extension),
                        redacted_text.encode("utf-8"),
                    )
                original_written = True

            if include_markdown and not (original_written and extension == "md"):
                _add_unique_output(
                    output_files,
                    used_output_names,
                    _build_output_name(filename, "_redacted", "md"),
                    redacted_text.encode("utf-8"),
                )
            if include_docx and not (original_written and extension == "docx"):
                _add_unique_output(
                    output_files,
                    used_output_names,
                    _build_output_name(filename, "_redacted", "docx"),
                    _build_docx(redacted_text),
                )
            continue

        if extension in {"png", "jpg", "jpeg"}:
            redacted_image = _redact_image(
                file_state["bytes"],
                selected_hits,
                categories,
                custom_words,
                engine_preference,
            )
            if include_original:
                _add_unique_output(
                    output_files,
                    used_output_names,
                    _build_output_name(filename, "_redacted", "png"),
                    redacted_image,
                )

            source_text = str(file_state.get("text", ""))
            redacted_text = apply_text_redaction(source_text, selected_hits, style=redaction_style)

            if include_markdown:
                _add_unique_output(
                    output_files,
                    used_output_names,
                    _build_output_name(filename, "_redacted", "md"),
                    redacted_text.encode("utf-8"),
                )
            if include_docx:
                _add_unique_output(
                    output_files,
                    used_output_names,
                    _build_output_name(filename, "_redacted", "docx"),
                    _build_docx(redacted_text),
                )

    report_json = json.dumps(report, ensure_ascii=False, indent=2).encode("utf-8")
    report_md = _create_markdown_report(report).encode("utf-8")

    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for name, payload in output_files:
            archive.writestr(name, payload)
        archive.writestr("report.json", report_json)
        archive.writestr("report.md", report_md)

    return zip_buffer.getvalue(), report
